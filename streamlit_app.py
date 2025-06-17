
import streamlit as st
import streamlit.components.v1 as components

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import json
from datetime import datetime

from utils.database import DatabaseManager
from utils.ai_engine import AIEngine
from utils.adaptive_logic import AdaptiveEngine
from config.settings import APP_CONFIG

import random

st.set_page_config(
    page_title="AdaptLearn - AI-Powered Learning",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded"
)

if 'user_id' not in st.session_state:
    st.session_state.user_id = None
if 'current_topic' not in st.session_state:
    st.session_state.current_topic = None
if 'quiz_state' not in st.session_state:
    st.session_state.quiz_state = {}

@st.cache_resource
def initialize_components():
    db = DatabaseManager()
    ai_engine = AIEngine()
    adaptive_engine = AdaptiveEngine()
    return db, ai_engine, adaptive_engine

db, ai_engine, adaptive_engine = initialize_components()

# ---- UI helpers ----------------------------------------------------------

def show_user_login():
    st.subheader("👤 Student Login")
    users = db.get_all_users()
    user_names = [u["name"] for u in users]
    option = st.radio("Choose option:", ["Select existing student", "Create new student"])
    if option == "Select existing student" and user_names:
        sel = st.selectbox("Select your name:", user_names, key="login_select")
        if st.button("Login"):
            user_id = next(u["id"] for u in users if u["name"] == sel)
            st.session_state.user_id = user_id
            st.rerun()
    elif option == "Create new student":
        new_name = st.text_input("Enter your name:")
        style = st.selectbox("Preferred learning style:",
                             ["Visual", "Auditory", "Kinesthetic", "Mixed"])
        if st.button("Create Account") and new_name:
            uid = db.create_user(new_name, style)
            st.session_state.user_id = uid
            st.success(f"Welcome, {new_name}!")
            st.rerun()

def nav_menu():
    user = db.get_user(st.session_state.user_id)
    st.write(f"👋 Welcome, **{user['name']}**")
    st.write(f"Learning Style: {user['learning_style']}")
    page = st.selectbox("Choose section:",
                        ["📊 Dashboard", "📚 Learn", "🧩 Quiz", "📈 Progress", "⚙️ Settings"])
    if st.button("Logout"):
        st.session_state.clear()
        st.rerun()
    return page

# ---- Pages ---------------------------------------------------------------

def page_welcome():
    c1, c2 = st.columns([2, 1])
    with c1:
        st.header("Welcome to AdaptLearn!")
        st.markdown("""
        ### 🚀 Features
        - **Adaptive Difficulty** – questions scale with your mastery  
        - **Personalised Content** – matches your learning style  
        - **Analytics Dashboard** – visualise your progress  
        - **AI‑Generated** – dynamic explanations & quizzes  
        """)
        st.info("👈 Log in via the sidebar to begin")
    with c2:
        st.image("https://via.placeholder.com/300x200/4CAF50/FFFFFF?text=AI+Learning")

def page_dashboard():
    st.header("📊 Your Learning Dashboard")
    stats = db.get_user_stats(st.session_state.user_id)
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Total Questions", stats["total_questions"])
    c2.metric("Correct", stats["correct_answers"])
    c3.metric("Accuracy", f"{stats['correct_answers']/max(stats['total_questions'],1)*100:.1f}%")
    c4.metric("Current Level", stats["current_level"])

    # Daily streak counter
    streak_data = db.get_user_streak_data(st.session_state.user_id)
    streak_count = streak_data["streak_count"]
    last_login_date = streak_data["last_login_date"]
    today = datetime.now().date()

    if last_login_date != today:
        st.warning("You haven't logged in today! Make sure to log in daily to maintain your streak.")
    else:
        if streak_count == 1:
          st.success(f"🎉 Daily Streak: {streak_count} day")
        else:
          st.success(f"🎉 Daily Streak: {streak_count} days")

    # Charts
    c1, c2 = st.columns(2)
    with c1:
        st.subheader("Progress Over Time")
        data = db.get_user_progress_data(st.session_state.user_id)
        if data:
            df = pd.DataFrame(data)
            fig = px.line(df, x="date", y="accuracy", title="Accuracy by Day")
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("No data yet – take some quizzes!")

    with c2:
        st.subheader("Subject Performance")
        sub = db.get_subject_stats(st.session_state.user_id)
        if sub:
            df = pd.DataFrame(sub)
            fig = px.bar(df, x="subject", y="accuracy", title="Accuracy by Subject")
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("No quizzes taken yet.")

def page_learn():
    st.header("📚 Learning Hub")

    subjects = APP_CONFIG["subjects"]
    subj = st.selectbox("Subject", list(subjects.keys()))
    topic = st.selectbox("Topic", subjects[subj])

    # --- Webcam drag-and-drop interaction ---

    if subj == "Science" and topic == "Biology":
      components.html("""
      <!DOCTYPE html>
      <html>
      <head>
        <meta charset="UTF-8">
        <title>Genetics Drag and Drop</title>
        <style>
          body { font-family: sans-serif; margin: 0; padding: 20px; }
          .container { display: flex; flex-wrap: wrap; gap: 10px; align-items: center; }
          .flower, .zone, .child {
            width: 80px;
            height: 80px;
            display: flex;
            align-items: center;
            justify-content: center;
            font-weight: bold;
            color: white;
            border-radius: 8px;
            cursor: grab;
            user-select: none;
          }
          .flower { border: 2px solid black; }
          .zone { background: #eee; color: white; border: 2px dashed #aaa; }
          .label { width: 100px; font-weight: bold; }
          #children-container { display: flex; align-items: center; margin-top: 20px; gap: 10px; }
          #children { display: flex; gap: 10px; flex-wrap: wrap; }
        </style>
      </head>
      <body>

      <h3>Drag one flower into the Parent A slot and one flower into the Parent B slot to explore inheritance.</h3>

      <div class="container" id="flowers">
        <div class="flower" draggable="true" data-geno="PP" style="background: purple;">PP</div>
        <div class="flower" draggable="true" data-geno="PP" style="background: purple;">PP</div>
        <div class="flower" draggable="true" data-geno="Pp" style="background: orchid;">Pp</div>
        <div class="flower" draggable="true" data-geno="Pp" style="background: orchid;">Pp</div>
        <div class="flower" draggable="true" data-geno="pp" style="background: lightgray;">pp</div>
        <div class="flower" draggable="true" data-geno="pp" style="background: lightgray;">pp</div>
      </div>

      <br>

      <div class="container">
        <div class="label">Parents:</div>
        <div class="zone" id="parentA" ondragover="event.preventDefault()" ondrop="drop(event, 'A')" style="color: gray;">Parent A</div>
        <div class="zone" id="parentB" ondragover="event.preventDefault()" ondrop="drop(event, 'B')" style="color: gray;">Parent B</div>
      </div>

      <div id="children-container">
        <div class="label">Offspring:</div>
        <div id="children"></div>
      </div>

      <script>
      let parents = { A: null, B: null };
      const colors = { PP: 'purple', Pp: 'orchid', pp: 'gray' };

      document.querySelectorAll('.flower').forEach(f => {
        f.addEventListener('dragstart', (e) => {
          e.dataTransfer.setData('text/plain', f.dataset.geno);
        });
      });

      function drop(e, zone) {
        const geno = e.dataTransfer.getData('text/plain');
        parents[zone] = geno;

        const box = document.getElementById('parent' + zone);
        box.textContent = geno;
        box.style.background = colors[geno] || '#eee';

        if (parents.A && parents.B) showChildren();
      }

      function getAlleles(geno) {
        return geno === 'PP' ? ['P','P'] :
              geno === 'Pp' ? ['P','p'] :
              ['p','p'];
      }

      function showChildren() {
        const a = getAlleles(parents.A);
        const b = getAlleles(parents.B);
        const combos = [];

        for (let i=0; i<2; i++) {
          for (let j=0; j<2; j++) {
            let genes = [a[i], b[j]].sort().join('');
            combos.push(genes === 'PP' ? 'PP' : genes === 'Pp' ? 'Pp' : 'pp');
          }
        }

        const out = combos.map(g =>
          `<div class='child flower' style='background:${colors[g]}'>${g}</div>`
        ).join('');
        document.getElementById('children').innerHTML = out;
      }
      </script>

      </body>
      </html>
      """, height=520)


      
    elif subj == "Science" and topic == "Chemistry":

      st.subheader("🧪🖐️ Drag both H₂ to O₂ to React")

      components.html("""
      <!DOCTYPE html>
      <html>
      <head>
        <style>
          video, canvas {
            position: absolute;
            top: 0;
            left: 0;
          }
          #wrapper {
            position: relative;
            width: 640px;
            height: 480px;
          }
          .label {
            position: absolute;
            font-weight: bold;
            color: white;
            text-align: center;
            width: 80px;
            pointer-events: none;
          }
        </style>
      </head>
      <body>
      <div id="wrapper">
        <video id="video" width="640" height="480" autoplay playsinline muted></video>
        <canvas id="canvas" width="640" height="480"></canvas>
      </div>

      <script src="https://cdn.jsdelivr.net/npm/@mediapipe/hands/hands.min.js"></script>
      <script src="https://cdn.jsdelivr.net/npm/@mediapipe/camera_utils/camera_utils.js"></script>

      <script>
      const video = document.getElementById("video");
      const canvas = document.getElementById("canvas");
      const ctx = canvas.getContext("2d");

      let draggingBox = null;
      let reacted = false;

      // Initial state: 2 H2 and 1 O2
      let boxes = [
        { id: "H1", label: "H₂", x: 100, y: 150, size: 80, color: "red", type: "H" },
        { id: "H2", label: "H₂", x: 200, y: 300, size: 80, color: "red", type: "H" },
        { id: "O", label: "O₂", x: 400, y: 200, size: 80, color: "blue", type: "O" }
      ];

      function drawScene(landmarks) {
        ctx.drawImage(video, 0, 0, canvas.width, canvas.height);
        for (const box of boxes) {
          ctx.fillStyle = box.color;
          ctx.fillRect(box.x, box.y, box.size, box.size);

          // Draw label
          ctx.fillStyle = "white";
          ctx.font = "bold 20px sans-serif";
          ctx.textAlign = "center";
          ctx.fillText(box.label, box.x + box.size / 2, box.y + box.size / 2 + 8);
        }

        if (landmarks) {
          for (const point of landmarks) {
            ctx.beginPath();
            ctx.arc(point.x * canvas.width, point.y * canvas.height, 5, 0, 2 * Math.PI);
            ctx.fillStyle = "lime";
            ctx.fill();
          }
        }
      }

      function isPinching(landmarks) {
        const dx = landmarks[4].x - landmarks[8].x;
        const dy = landmarks[4].y - landmarks[8].y;
        return Math.sqrt(dx * dx + dy * dy) < 0.05;
      }

      function checkOverlap(boxA, boxB) {
        return (
          boxA.x < boxB.x + boxB.size &&
          boxA.x + boxA.size > boxB.x &&
          boxA.y < boxB.y + boxB.size &&
          boxA.y + boxA.size > boxB.y
        );
      }

      function handleReaction() {
        const hBoxes = boxes.filter(b => b.type === "H");
        const oBox = boxes.find(b => b.type === "O");

        if (!oBox || hBoxes.length < 2) return false;

        const overlap1 = checkOverlap(hBoxes[0], oBox);
        const overlap2 = checkOverlap(hBoxes[1], oBox);

        return overlap1 && overlap2;
      }

      function reactToWater() {
        boxes = [
          { id: "H2O1", label: "H₂O", x: 200, y: 150, size: 100, color: "green", type: "H2O" },
          { id: "H2O2", label: "H₂O", x: 300, y: 250, size: 100, color: "green", type: "H2O" }
        ];
        reacted = true;
      }

      function updateDragging(landmarks) {
        const x = landmarks[8].x * canvas.width;
        const y = landmarks[8].y * canvas.height;

        if (draggingBox && isPinching(landmarks)) {
          draggingBox.x = x - draggingBox.size / 2;
          draggingBox.y = y - draggingBox.size / 2;
        } else {
          draggingBox = null;
          if (!reacted && isPinching(landmarks)) {
            for (let box of boxes) {
              if (
                x >= box.x && x <= box.x + box.size &&
                y >= box.y && y <= box.y + box.size
              ) {
                draggingBox = box;
                break;
              }
            }
          }
        }

        if (!reacted && handleReaction()) {
          reactToWater();
        }
      }

      const hands = new Hands({
        locateFile: (file) => `https://cdn.jsdelivr.net/npm/@mediapipe/hands/${file}`,
      });

      hands.setOptions({
        maxNumHands: 1,
        modelComplexity: 1,
        minDetectionConfidence: 0.7,
        minTrackingConfidence: 0.5,
      });

      hands.onResults((results) => {
        if (results.multiHandLandmarks.length > 0) {
          const landmarks = results.multiHandLandmarks[0];
          updateDragging(landmarks);
          drawScene(landmarks);
        } else {
          drawScene(null);
          draggingBox = null;
        }
      });

      const camera = new Camera(video, {
        onFrame: async () => await hands.send({ image: video }),
        width: 640,
        height: 480,
      });
      camera.start();
      </script>
      </body>
      </html>
      """, height=500)


    # --- Content generation from subject/topic ---
    # st.subheader("🎯 Personalised Learning Content")
    # if st.button("Generate Content"):
    #     lvl = db.get_user_topic_level(st.session_state.user_id, subj, topic)
    #     with st.spinner("Generating personalised lesson…"):
    #         content = ai_engine.generate_learning_content(subj, topic, lvl)
    #     st.markdown(content, unsafe_allow_html=True)
    #     if st.button("Take Quiz"):
    #         st.session_state.quiz_topic = {"subject": subj, "topic": topic}
    #         st.switch_page("🧩 Quiz")

    # --- File Upload & Quiz Generation ---
    # st.subheader("📄 Upload Content for Quiz Generation")
    # uploaded_file = st.file_uploader("Upload a text file or document", type=["txt", "pdf", "docx"])

    # if uploaded_file:
    #     def extract_file_content(file):
    #         if file.name.endswith(".txt"):
    #             return file.read().decode("utf-8")
    #         elif file.name.endswith(".pdf"):
    #             import PyPDF2
    #             reader = PyPDF2.PdfReader(file)
    #             return "".join(page.extract_text() for page in reader.pages)
    #         elif file.name.endswith(".docx"):
    #             import docx
    #             doc = docx.Document(file)
    #             return "\n".join(paragraph.text for paragraph in doc.paragraphs)
    #         else:
    #             raise ValueError("Unsupported file format.")

    #     with st.spinner("Processing file..."):
    #         content = extract_file_content(uploaded_file)
    #     st.success("File uploaded successfully!")

    #     if st.button("Generate Quiz"):
    #         quiz_questions = ai_engine.generate_quiz_questions(content)
    #         st.session_state.quiz_questions = quiz_questions
    #         st.session_state.quiz_attempts = []
    #         st.success("Quiz generated! Navigate to the Quiz page to start.")

    # # --- Retest Functionality ---
    # if st.session_state.get("quiz_attempts"):
    #     st.subheader("🔁 Retest on Incorrect Questions")
    #     if st.button("Retest"):
    #         incorrect_questions = [
    #             attempt["question"] for attempt in st.session_state.quiz_attempts if not attempt["is_correct"]
    #         ]
    #         if incorrect_questions:
    #             retest_questions = ai_engine.generate_similar_questions(incorrect_questions)
    #             st.session_state.quiz_questions = retest_questions
    #             st.success("Retest questions generated! Navigate to the Quiz page to start.")
    #         else:
    #             st.info("No incorrect questions to retest.")


def page_quiz():
    subjects = APP_CONFIG["subjects"]

    # If quiz hasn't started, show selection
    if "quiz_topic" not in st.session_state:
        subj = st.selectbox("Subject", list(subjects.keys()), key="quiz_subject_select")
        topic = st.selectbox("Topic", subjects[subj], key="quiz_topic_select")

        if st.button("Start Quiz"):
            st.session_state.quiz_topic = {"subject": subj, "topic": topic}
            st.session_state.quiz_topic_prev = {"subject": subj, "topic": topic}
            st.session_state.quiz_state = {}
            st.rerun()
        return  # Exit early until quiz starts

    # Quiz has started
    info = st.session_state.get("quiz_topic")
    if not info:
        st.info("Click 'Start Quiz' to begin.")
        return

    subj, topic = info["subject"], info["topic"]

    # Reset quiz state if subject/topic changed
    if (
        "quiz_topic_prev" not in st.session_state
        or st.session_state.quiz_topic != st.session_state.quiz_topic_prev
    ):
        st.session_state.quiz_state = {}
        st.session_state.quiz_topic_prev = st.session_state.quiz_topic

    level = db.get_user_topic_level(st.session_state.user_id, subj, topic)
    db.ensure_user_topic_entry(st.session_state.user_id, subj, topic, level)

    qs = st.session_state.quiz_state
    if not qs:
        st.session_state.quiz_state = {"n": 0, "correct": 0, "level": level, "q": None}
        qs = st.session_state.quiz_state

    # Hardcoded fallback question bank
    fallback_bank = {
        "Science": {
            "Chemistry": {
                1: [
                    {
                        "question": "Which of the following is a combustion reaction?",
                        "options": ["CH₄ + 2O₂ → CO₂ + 2H₂O", "HCl + NaOH → NaCl + H₂O", "AgNO₃ + NaCl → AgCl + NaNO₃", "Zn + HCl → ZnCl₂ + H₂"],
                        "correct_answer": "CH₄ + 2O₂ → CO₂ + 2H₂O",
                        "explanation": "Combustion reactions involve a substance reacting with oxygen to produce heat and light."
                    },
                    {
                        "question": "Which type of chemical reaction is represented by: 2H₂O₂ → 2H₂O + O₂?",
                        "options": ["Synthesis", "Decomposition", "Single displacement", "Combustion"],
                        "correct_answer": "Decomposition",
                        "explanation": "Hydrogen peroxide breaks down into water and oxygen — a classic decomposition reaction."
                    },
                    {
                        "question": "What do all acids release in aqueous solution?",
                        "options": ["OH⁻ ions", "H⁺ ions", "Na⁺ ions", "Cl⁻ ions"],
                        "correct_answer": "H⁺ ions",
                        "explanation": "Acids release hydrogen (H⁺) ions when dissolved in water."
                    },
                    {
                        "question": "Which substance is an element?",
                        "options": ["H₂", "HCl", "NaCl", "H₂O"],
                        "correct_answer": "H₂",
                        "explanation": "H₂ is a pure element made of two hydrogen atoms."
                    },
                    {
                        "question": "Which of the following is a physical change?",
                        "options": ["Boiling water", "Burning paper", "Rusting iron", "Baking a cake"],
                        "correct_answer": "Boiling water",
                        "explanation": "Boiling is a physical change as no new substance is formed."
                    },
                ],
                2: [
                    {
                        "question": "What type of chemical reaction is: 2KClO₃ → 2KCl + 3O₂?",
                        "options": ["Synthesis", "Decomposition", "Single Displacement", "Combustion"],
                        "correct_answer": "Decomposition",
                        "explanation": "One compound breaks down into simpler substances, characteristic of decomposition."
                    },
                    {
                        "question": "Which of the following is a sign that a chemical reaction has occurred?",
                        "options": ["Melting of ice", "Color change", "Breaking glass", "Boiling water"],
                        "correct_answer": "Color change",
                        "explanation": "Color change often indicates a new substance has formed during a chemical reaction."
                    },
                    {
                        "question": "Which metal is most reactive?",
                        "options": ["Gold", "Iron", "Potassium", "Copper"],
                        "correct_answer": "Potassium",
                        "explanation": "Potassium is highly reactive, especially with water."
                    },
                    {
                        "question": "What is the pH of a neutral solution?",
                        "options": ["7", "0", "14", "10"],
                        "correct_answer": "7",
                        "explanation": "A pH of 7 indicates a neutral solution like pure water."
                    },
                    {
                        "question": "Which gas is usually released during a metal-acid reaction?",
                        "options": ["Oxygen", "Carbon dioxide", "Hydrogen", "Nitrogen"],
                        "correct_answer": "Hydrogen",
                        "explanation": "Metals reacting with acids release hydrogen gas."
                    },
                ],
                3: [
                    {
                        "question": "Which of the following best represents a redox reaction?",
                        "options": ["H₂SO₄ + NaOH → Na₂SO₄ + H₂O", "AgNO₃ + NaCl → AgCl + NaNO₃", "2Fe + 3Cl₂ → 2FeCl₃", "CaCO₃ → CaO + CO₂"],
                        "correct_answer": "2Fe + 3Cl₂ → 2FeCl₃",
                        "explanation": "Iron is oxidized and chlorine is reduced — both oxidation and reduction occur."
                    },
                    {
                        "question": "Which process involves both oxidation and reduction?",
                        "options": ["Redox reaction", "Precipitation", "Neutralization", "Hydrolysis"],
                        "correct_answer": "Redox reaction",
                        "explanation": "Redox reactions involve simultaneous oxidation and reduction."
                    },
                    {
                        "question": "What is the oxidation number of hydrogen in most compounds?",
                        "options": ["+1", "0", "-1", "+2"],
                        "correct_answer": "+1",
                        "explanation": "Hydrogen usually has a +1 oxidation number in compounds."
                    },
                    {
                        "question": "Which indicator turns red in acid and blue in alkali?",
                        "options": ["Phenolphthalein", "Methyl orange", "Litmus", "Universal Indicator"],
                        "correct_answer": "Litmus",
                        "explanation": "Litmus is a common indicator: red in acid, blue in base."
                    },
                    {
                        "question": "Which element has the highest electronegativity?",
                        "options": ["Oxygen", "Chlorine", "Fluorine", "Nitrogen"],
                        "correct_answer": "Fluorine",
                        "explanation": "Fluorine is the most electronegative element."
                    },
                ]
            },
            "Biology": {
                1: [
                    {
                        "question": "What is the basic unit of heredity?",
                        "options": ["Cell", "Gene", "Chromosome", "Protein"],
                        "correct_answer": "Gene",
                        "explanation": "Genes are segments of DNA that carry hereditary information from parents to offspring."
                    },
                    {
                        "question": "Which of these is an example of a dominant trait?",
                        "options": ["Blue eyes", "Attached earlobes", "Widow's peak", "Straight hairline"],
                        "correct_answer": "Widow's peak",
                        "explanation": "Widow’s peak is a common example of a dominant trait passed from one generation to the next."
                    },
                    {
                        "question": "Where are genes located?",
                        "options": ["Cytoplasm", "Ribosomes", "Chromosomes", "Nucleus membrane"],
                        "correct_answer": "Chromosomes",
                        "explanation": "Genes are segments of DNA on chromosomes in the nucleus."
                    },
                    {
                        "question": "What does DNA stand for?",
                        "options": ["Deoxyribonucleic acid", "Deoxyribonuclear acid", "Dinucleic acid", "Dual nucleic acid"],
                        "correct_answer": "Deoxyribonucleic acid",
                        "explanation": "DNA is short for Deoxyribonucleic acid."
                    },
                    {
                        "question": "What determines the trait expressed by a gene?",
                        "options": ["The RNA sequence", "The environment", "The allele pair", "The cytoplasm"],
                        "correct_answer": "The allele pair",
                        "explanation": "Traits are determined by the combination of alleles (dominant/recessive)."
                    },
                ],
                2: [
                    {
                        "question": "In Mendel’s experiments, what was the ratio of dominant to recessive traits in the F2 generation?",
                        "options": ["1:1", "3:1", "9:3:3:1", "2:1"],
                        "correct_answer": "3:1",
                        "explanation": "In the F2 generation, Mendel observed a 3:1 ratio of dominant to recessive traits."
                    },
                    {
                        "question": "Which genotype represents a heterozygous individual?",
                        "options": ["AA", "aa", "Aa", "BB"],
                        "correct_answer": "Aa",
                        "explanation": "Heterozygous individuals have two different alleles for a trait, such as 'A' and 'a'."
                    },
                    {
                        "question": "What is a phenotype?",
                        "options": ["Genetic makeup", "Physical expression of traits", "Allele pair", "DNA sequence"],
                        "correct_answer": "Physical expression of traits",
                        "explanation": "Phenotype is how a trait appears based on genotype."
                    },
                    {
                        "question": "Which pair of chromosomes determines human sex?",
                        "options": ["1 and 2", "11 and 12", "X and Y", "A and B"],
                        "correct_answer": "X and Y",
                        "explanation": "The X and Y chromosomes determine biological sex."
                    },
                    {
                        "question": "Which of the following is *not* inherited genetically?",
                        "options": ["Eye color", "Blood type", "Language spoken", "Hair color"],
                        "correct_answer": "Language spoken",
                        "explanation": "Language is learned, not inherited through genes."
                    },
                ],
                3: [
                    {
                        "question": "A woman with blood type AB and a man with blood type O have a child. What are the possible blood types of the child?",
                        "options": ["A, B, AB, or O", "A or B", "AB only", "O only"],
                        "correct_answer": "A or B",
                        "explanation": "The AB parent contributes either A or B; the O parent contributes only O. So the child can be A or B."
                    },
                    {
                        "question": "In a dihybrid cross between two heterozygous individuals (RrYy × RrYy), what fraction of the offspring will show both recessive traits?",
                        "options": ["1/4", "1/8", "1/16", "9/16"],
                        "correct_answer": "1/16",
                        "explanation": "Only the offspring with the genotype rryy will express both recessive traits: 1 out of 16 combinations."
                    },
                    {
                        "question": "What is codominance?",
                        "options": ["One allele is dominant", "Both alleles are recessive", "Both alleles are expressed equally", "Neither allele is expressed"],
                        "correct_answer": "Both alleles are expressed equally",
                        "explanation": "In codominance, both traits appear in the phenotype."
                    },
                    {
                        "question": "Which disorder is caused by a single gene mutation?",
                        "options": ["Cancer", "Sickle Cell Anemia", "Diabetes", "Autism"],
                        "correct_answer": "Sickle Cell Anemia",
                        "explanation": "Sickle Cell is a well-known example of a single-gene mutation."
                    },
                    {
                        "question": "What is the chance of a carrier father and carrier mother having an affected child with a recessive condition?",
                        "options": ["25%", "50%", "75%", "100%"],
                        "correct_answer": "25%",
                        "explanation": "With both carriers, the chance is 1 in 4 (25%) for a recessive condition."
                    },
                ]
            }
        }
    }

    # Prevent displaying questions after 5
    if qs["n"] >= 5:
        if st.button("Finish"):
            acc = qs["correct"] / qs["n"] * 100
            st.success(f"🎉 You answered {qs['correct']} out of {qs['n']} correctly ({acc:.1f}%)!")

            with st.expander("📘 Review your answers"):
                for i, entry in enumerate(qs.get("history", []), 1):
                    st.markdown(f"**Q{i}:** {entry['question']}")
                    st.markdown(f"- Your answer: {entry['your_answer']}")
                    st.markdown(f"- Correct answer: {entry['correct_answer']}")
                    if entry["explanation"]:
                        st.markdown(f"_Explanation:_ {entry['explanation']}")
                    st.markdown("---")

            st.session_state.quiz_state = {}
        return

    # Get next question
    if qs["q"] is None:
        bank = fallback_bank.get(subj, {}).get(topic, {})
        level_questions = bank.get(min(qs["level"], 3), [])
        if qs["n"] < len(level_questions):
            qs["q"] = level_questions[qs["n"]]
        elif level_questions:
            qs["q"] = random.choice(level_questions)
        else:
            st.error("No questions available for this topic.")
            return

    # Show question
    q = qs["q"]
    st.subheader(f"{subj} – {topic}")
    st.write(f"**Q{qs['n'] + 1}:** {q['question']}")
    ans = st.radio("Answer:", q["options"], key=f"q_{qs['n']}")

    if st.button("Submit"):
        correct = (ans == q["correct_answer"])

        qs.setdefault("history", []).append({
            "question": q["question"],
            "your_answer": ans,
            "correct_answer": q["correct_answer"],
            "explanation": q.get("explanation", ""),
            "correct": correct
        })

        qs["n"] += 1
        if correct:
            qs["correct"] += 1

        db.record_quiz_answer(
            st.session_state.user_id,
            subj, topic,
            q["question"], ans, q["correct_answer"], correct,
            qs["level"]
        )
        qs["level"] = adaptive_engine.adjust_difficulty(qs["level"], correct, qs["n"], qs["correct"])
        qs["q"] = None
        st.rerun()
        
        
def page_progress():
    st.header("📈 Detailed Progress Analytics")
    stats = db.get_user_stats(st.session_state.user_id)
    acc = stats['correct_answers']/max(stats['total_questions'],1)*100
    fig = go.Figure(go.Indicator(mode="gauge+number", value=acc,
                                 gauge={'axis':{'range':[0,100]}}))
    st.plotly_chart(fig, use_container_width=True)

def page_settings():
    st.header("⚙️ Settings")
    user = db.get_user(st.session_state.user_id)
    with st.form("settings"):
        name = st.text_input("Name", user["name"])
        style = st.selectbox("Learning Style", APP_CONFIG["learning_styles"],
                             index=APP_CONFIG["learning_styles"].index(user["learning_style"]))
        if st.form_submit_button("Save"):
            db.update_user_settings(st.session_state.user_id, {"name": name, "learning_style": style})
            st.success("Saved!")
            st.rerun()

def extract_file_content(uploaded_file) -> str:
    """
    Extracts text content from the uploaded file.
    Supports .txt, .pdf, and .docx formats.
    """
    try:
        if uploaded_file.name.endswith(".txt"):
            return uploaded_file.read().decode("utf-8")
        elif uploaded_file.name.endswith(".pdf"):
            import PyPDF2
            reader = PyPDF2.PdfReader(uploaded_file)
            return "".join(page.extract_text() for page in reader.pages)
        elif uploaded_file.name.endswith(".docx"):
            import docx
            doc = docx.Document(uploaded_file)
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)
        else:
            raise ValueError("Unsupported file format.")
    except Exception as e:
        raise ValueError(f"Error processing file: {e}")
    
# ---- Main ---------------------------------------------------------------

def main():
    with st.sidebar:
        if st.session_state.user_id is None:
            show_user_login()
        else:
            page = nav_menu()

    if st.session_state.user_id is None:
        page_welcome()
    else:
        if page == "📊 Dashboard":
            page_dashboard()
        elif page == "📚 Learn":
            page_learn()
        elif page == "🧩 Quiz":
            page_quiz()
        elif page == "📈 Progress":
            page_progress()
        elif page == "⚙️ Settings":
            page_settings()

if __name__ == "__main__":
    main()
